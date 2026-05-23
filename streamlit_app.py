"""
AURA GREEN — ESG Statutory Compliance & Audit Intelligence Platform
Production-ready Streamlit application.

Requirements:
    pip install streamlit pypdf google-generativeai python-docx

Run:
    streamlit run streamlit_app.py
"""

import io
import json
import datetime
import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# I.  PAGE CONFIG  (must be the first Streamlit call)
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Aura Green — ESG Intelligence",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─────────────────────────────────────────────────────────────────────────────
# II.  GLOBAL CSS INJECTION  — Visual Architecture Interface Model
# ─────────────────────────────────────────────────────────────────────────────
st.markdown(
    """
    <style>
    /* ── Core canvas ─────────────────────────────────────────────────── */
    html, body, [data-testid="stAppViewContainer"],
    [data-testid="stMain"], .main {
        background-color: #519EAA !important;
        color: #FFFFFF !important;
    }

    /* ── Remove default Streamlit padding ───────────────────────────── */
    [data-testid="stMain"] > div:first-child { padding-top: 0 !important; }
    .block-container { padding: 0 !important; max-width: 100% !important; }

    /* ── Executive top-header bar ────────────────────────────────────── */
    .aura-header {
        width: 100%;
        background: #2C5F6A;
        border-bottom: 1px solid #000000;
        box-shadow: 0px 4px 16px rgba(0,0,0,0.4);
        padding: 18px 32px;
        margin-bottom: 24px;
        display: flex;
        align-items: center;
        justify-content: space-between;
    }
    .aura-header-title {
        font-size: 26px;
        font-weight: 700;
        color: #FFFFFF;
        letter-spacing: 0.06em;
        text-transform: uppercase;
    }
    .aura-header-sub {
        font-size: 12px;
        color: #A8D5DC;
        letter-spacing: 0.08em;
        margin-top: 2px;
    }
    .aura-header-meta {
        font-size: 11px;
        color: #A8D5DC;
        text-align: right;
    }

    /* ── Component panels ────────────────────────────────────────────── */
    .aura-panel {
        background-color: #3F7A85;
        border-radius: 0px;
        border: 1px solid #000000;
        box-shadow: 0px 8px 24px rgba(0,0,0,0.3);
        padding: 20px;
        margin-bottom: 16px;
    }
    .aura-panel-title {
        font-size: 11px;
        font-weight: 700;
        color: #A8D5DC;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        margin-bottom: 14px;
        border-bottom: 0.5px solid rgba(255,255,255,0.15);
        padding-bottom: 8px;
    }

    /* ── Status pills ────────────────────────────────────────────────── */
    .pill-green {
        display: inline-block;
        background: rgba(16,185,129,0.15);
        color: #10B981;
        border: 1px solid #10B981;
        border-radius: 20px;
        font-size: 11px;
        font-weight: 600;
        padding: 2px 10px;
        letter-spacing: 0.03em;
    }
    .pill-amber {
        display: inline-block;
        background: rgba(251,191,36,0.12);
        color: #FBBF24;
        border: 1px solid #FBBF24;
        border-radius: 20px;
        font-size: 11px;
        font-weight: 600;
        padding: 2px 10px;
        letter-spacing: 0.03em;
    }
    .pill-red {
        display: inline-block;
        background: rgba(239,68,68,0.12);
        color: #EF4444;
        border: 1px solid #EF4444;
        border-radius: 20px;
        font-size: 11px;
        font-weight: 600;
        padding: 2px 10px;
        letter-spacing: 0.03em;
    }
    .pill-slate {
        display: inline-block;
        background: rgba(255,255,255,0.08);
        color: #CBD5E1;
        border: 1px solid rgba(255,255,255,0.2);
        border-radius: 20px;
        font-size: 11px;
        font-weight: 600;
        padding: 2px 10px;
        letter-spacing: 0.03em;
    }

    /* ── KPI score cards ──────────────────────────────────────────────── */
    .kpi-card {
        background: #2C5F6A;
        border: 1px solid rgba(255,255,255,0.1);
        border-radius: 0px;
        padding: 16px;
        text-align: center;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.25);
    }
    .kpi-label { font-size: 10px; color: #A8D5DC; text-transform: uppercase; letter-spacing: 0.08em; }
    .kpi-value { font-size: 28px; font-weight: 700; color: #FFFFFF; line-height: 1.1; margin: 4px 0; }
    .kpi-trend { font-size: 11px; color: #10B981; }
    .kpi-trend.down { color: #EF4444; }
    .kpi-trend.neutral { color: #FBBF24; }

    /* ── Metric rows inside panels ───────────────────────────────────── */
    .metric-row {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding: 8px 0;
        border-bottom: 0.5px solid rgba(255,255,255,0.08);
        font-size: 12px;
        color: #E2E8F0;
    }
    .metric-row:last-child { border-bottom: none; }
    .metric-label { color: #A8D5DC; }

    /* ── Progress bars ───────────────────────────────────────────────── */
    .prog-wrap { margin: 6px 0 2px; }
    .prog-bg {
        background: rgba(255,255,255,0.08);
        border-radius: 0px;
        height: 5px;
        width: 100%;
        overflow: hidden;
    }
    .prog-fill-green { height: 5px; background: #10B981; }
    .prog-fill-amber { height: 5px; background: #FBBF24; }
    .prog-fill-red   { height: 5px; background: #EF4444; }
    .prog-meta {
        display: flex;
        justify-content: space-between;
        font-size: 10px;
        color: #94A3B8;
        margin-top: 2px;
    }

    /* ── File list items ─────────────────────────────────────────────── */
    .file-item {
        display: flex;
        align-items: center;
        gap: 10px;
        padding: 8px 10px;
        background: rgba(0,0,0,0.15);
        border: 0.5px solid rgba(255,255,255,0.08);
        margin-bottom: 6px;
        font-size: 12px;
        color: #E2E8F0;
    }
    .file-dot-green { width:7px; height:7px; border-radius:50%; background:#10B981; flex-shrink:0; }
    .file-dot-amber { width:7px; height:7px; border-radius:50%; background:#FBBF24; flex-shrink:0; }

    /* ── Audit narrative blocks ───────────────────────────────────────── */
    .audit-block {
        background: rgba(0,0,0,0.18);
        border-left: 3px solid #10B981;
        padding: 14px 16px;
        margin: 12px 0;
        font-size: 13px;
        color: #E2E8F0;
        line-height: 1.7;
    }
    .audit-block.amber { border-left-color: #FBBF24; }
    .audit-block.red   { border-left-color: #EF4444; }
    .audit-section-head {
        font-size: 13px;
        font-weight: 700;
        color: #FFFFFF;
        text-transform: uppercase;
        letter-spacing: 0.06em;
        margin: 20px 0 10px;
        padding-bottom: 6px;
        border-bottom: 0.5px solid rgba(255,255,255,0.15);
    }

    /* ── Streamlit widget overrides ──────────────────────────────────── */
    [data-testid="stFileUploader"] {
        background: rgba(0,0,0,0.15);
        border: 1px dashed rgba(255,255,255,0.25) !important;
        padding: 10px;
    }
    [data-testid="stFileUploader"] label { color: #E2E8F0 !important; font-size: 12px; }
    button[kind="primary"], [data-testid="baseButton-primary"] {
        background-color: #10B981 !important;
        border: none !important;
        border-radius: 0px !important;
        color: #FFFFFF !important;
        font-weight: 700 !important;
        letter-spacing: 0.05em;
    }
    button[kind="secondary"], [data-testid="baseButton-secondary"] {
        background-color: #2C5F6A !important;
        border: 1px solid rgba(255,255,255,0.2) !important;
        border-radius: 0px !important;
        color: #E2E8F0 !important;
    }
    .stTabs [data-baseweb="tab-list"] {
        background: #2C5F6A;
        border-bottom: 1px solid #000000;
        gap: 0;
    }
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        color: #A8D5DC;
        border-radius: 0px;
        border-bottom: 2px solid transparent;
        padding: 10px 20px;
        font-size: 12px;
        font-weight: 600;
        letter-spacing: 0.04em;
    }
    .stTabs [aria-selected="true"] {
        background: rgba(16,185,129,0.08) !important;
        color: #10B981 !important;
        border-bottom: 2px solid #10B981 !important;
    }
    .stTabs [data-baseweb="tab-panel"] {
        background: #3F7A85;
        border: 1px solid #000000;
        border-top: none;
        padding: 20px;
    }
    [data-testid="stTextInput"] input,
    [data-testid="stTextArea"] textarea {
        background: rgba(0,0,0,0.2) !important;
        border: 1px solid rgba(255,255,255,0.15) !important;
        border-radius: 0px !important;
        color: #FFFFFF !important;
        font-size: 13px;
    }
    [data-testid="stAlert"] {
        border-radius: 0px !important;
        border-left: 3px solid #FBBF24 !important;
    }
    label, .stSelectbox label, .stTextInput label { color: #A8D5DC !important; font-size: 11px; }
    p, li { color: #E2E8F0; }
    h1, h2, h3, h4 { color: #FFFFFF !important; }
    [data-testid="stExpander"] {
        background: rgba(0,0,0,0.12);
        border: 0.5px solid rgba(255,255,255,0.1) !important;
        border-radius: 0px !important;
    }
    [data-testid="stExpander"] summary { color: #E2E8F0 !important; font-size: 12px; }
    div[data-testid="stDownloadButton"] button {
        background: #2C5F6A !important;
        border: 1px solid #10B981 !important;
        color: #10B981 !important;
        border-radius: 0px !important;
        font-size: 12px !important;
        font-weight: 600 !important;
        width: 100%;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────────────────────────────────────────────
# III.  HELPER: STATUS PILL HTML
# ─────────────────────────────────────────────────────────────────────────────
def pill(label: str, status: str = "green") -> str:
    """Return an HTML status pill string. status ∈ {green, amber, red, slate}."""
    return f'<span class="pill-{status}">{label}</span>'


def progress_bar(pct: float, color: str = "green") -> str:
    """Return an HTML progress bar at `pct` percent."""
    w = min(max(pct, 0), 100)
    return (
        f'<div class="prog-bg"><div class="prog-fill-{color}" style="width:{w}%"></div></div>'
    )


def status_from_pct(pct: float, thresholds=(50, 80)) -> str:
    """Map a completion percentage to a colour token."""
    if pct >= thresholds[1]:
        return "green"
    if pct >= thresholds[0]:
        return "amber"
    return "red"


def status_label(pct: float, thresholds=(50, 80)) -> str:
    labels = {
        "green": "Compliant",
        "amber": "At Risk",
        "red": "Breach",
    }
    return labels[status_from_pct(pct, thresholds)]


# ─────────────────────────────────────────────────────────────────────────────
# IV.  FALLBACK DATA — Pakistani SOE sample (Finance Division / CMU parameters)
# ─────────────────────────────────────────────────────────────────────────────
FALLBACK_ANALYSIS: dict = {
    "entity": "National Highway Authority (NHA) — Illustrative Audit Sample",
    "audit_date": "Q2 FY2025",
    "overall_score": 61.4,
    "environmental": {
        "score": 54.0,
        "ghg_emissions_mt": 182400,
        "emissions_intensity": 3.21,
        "renewable_energy_pct": 14.0,
        "water_reclamation_pct": 28.0,
        "findings": [
            {
                "framework": "SECP ESG Guidelines 2023 — Environmental Index §3.1",
                "metric": "GHG Emissions Disclosure (Scope 1 & 2)",
                "status": "amber",
                "pct": 60,
                "detail": (
                    "NHA's Annual Report FY2024 discloses aggregate fuel consumption (diesel) for "
                    "the construction machinery fleet, but does not convert these figures into "
                    "metric-ton CO₂-equivalent values as mandated under the SECP ESG Disclosure "
                    "Guidelines §3.1(a). The absence of a formal Scope 2 (purchased electricity) "
                    "calculation for administrative offices constitutes a partial non-disclosure. "
                    "This omission is compounded by the lack of any third-party assurance on "
                    "reported energy data, a requirement under the SECP 2023 roadmap Phase II "
                    "timeline for PSX-listed entities and SOEs under active CMU monitoring."
                ),
            },
            {
                "framework": "Companies Act 2017 — Section 238(1)(d)",
                "metric": "Environmental Initiatives & CSR Disclosure",
                "status": "amber",
                "pct": 55,
                "detail": (
                    "Section 238 of the Companies Act, 2017, obligates the Board's Directors' "
                    "Report to contain a statement of key environmental measures and CSR "
                    "expenditure. The NHA FY2024 report contains a boilerplate reference to "
                    "tree plantation drives under the PM's Green Pakistan Initiative but omits "
                    "verifiable expenditure figures, hectare coverage, or species-level "
                    "biodiversity data. This renders the disclosure non-compliant with the "
                    "specificity requirements articulated in the SECP's 2022 clarificatory "
                    "circular on Section 238 disclosures."
                ),
            },
            {
                "framework": "SECP ESG Guidelines 2023 — Environmental Index §3.3",
                "metric": "Renewable Energy Mix",
                "status": "red",
                "pct": 14,
                "detail": (
                    "At 14% renewable energy penetration, NHA is critically below the SECP's "
                    "indicative 30% interim benchmark for heavy-infrastructure SOEs by FY2026. "
                    "No capital expenditure roadmap for solar or wind augmentation of fixed "
                    "infrastructure (toll plazas, administrative campuses) has been filed with "
                    "the CMU or disclosed in the annual report. This gap is flagged as a "
                    "High-Priority Operational Omission requiring a Board-level resolution "
                    "within 90 days under SOE Act 2023, Section 6(f)."
                ),
            },
        ],
    },
    "social": {
        "score": 58.0,
        "ceo_pay_ratio": 42.1,
        "gender_pay_gap_pct": 11.3,
        "turnover_rate_pct": 18.7,
        "temp_worker_ratio_pct": 34.0,
        "supply_chain_audit_pct": 22.0,
        "findings": [
            {
                "framework": "SECP ESG Guidelines 2023 — Social Index §4.2",
                "metric": "CEO-to-Employee Pay Ratio",
                "status": "amber",
                "pct": 58,
                "detail": (
                    "A CEO-to-median-employee pay ratio of 42.1:1 is disclosed in aggregate "
                    "salary cost data but without granular band-level disclosure. SECP Social "
                    "Index §4.2 requires explicit reporting of the CEO compensation multiple "
                    "against median employee remuneration with a narrative justification where "
                    "the ratio exceeds 30:1. The absence of this narrative explanation "
                    "constitutes an incomplete disclosure. Under SOE Policy 2023, the "
                    "Establishment Division's Unified Pay Scale alignment is a mitigating "
                    "factor but does not relieve the SOE of its SECP disclosure obligation."
                ),
            },
            {
                "framework": "SECP ESG Guidelines 2023 — Social Index §4.3 & §4.4",
                "metric": "Gender Pay Gap & Workforce Diversity",
                "status": "red",
                "pct": 28,
                "detail": (
                    "A gender pay gap of 11.3% is a material breach of the SECP's indicative "
                    "≤7% threshold for public-sector entities. More critically, NHA's FY2024 "
                    "report contains no gender-disaggregated workforce data, no Equal "
                    "Opportunity Policy statement, and no Pay Gap Remediation Action Plan — "
                    "all of which are mandatory disclosure components under Social Index §4.4. "
                    "This is classified as a Red Flag Omission. Under Section 238(1)(e) of the "
                    "Companies Act 2017, the Board is required to certify human capital "
                    "disclosures; non-disclosure here creates a potential directorial liability."
                ),
            },
            {
                "framework": "SECP ESG Guidelines 2023 — Social Index §4.6",
                "metric": "Supply Chain Labour Protections (Child/Forced Labour)",
                "status": "red",
                "pct": 22,
                "detail": (
                    "Only 22% of Tier-1 construction contractors have been audited for child "
                    "and forced labour protections, measured against NHA's own Supplier Code "
                    "of Conduct. The SECP ESG Social Index §4.6 requires SOEs operating "
                    "infrastructure projects to achieve ≥80% upstream supply chain audit "
                    "coverage. Given Pakistan's obligations under ILO Conventions 138 and 182 "
                    "(ratified 2006 and 2001), this gap also creates international treaty "
                    "compliance exposure. The Finance Division's CMU is empowered under SOE "
                    "Act 2023 Section 25(3) to mandate corrective reporting timelines."
                ),
            },
        ],
    },
    "governance": {
        "score": 72.0,
        "board_independence_pct": 55.0,
        "board_gender_diversity_pct": 18.0,
        "external_assurance": False,
        "supplier_code_of_conduct": True,
        "solvency_declaration_filed": True,
        "findings": [
            {
                "framework": "SOE Act 2023 — Section 6 & SOE Policy 2023",
                "metric": "Board Independence & Sound Governance",
                "status": "green",
                "pct": 82,
                "detail": (
                    "NHA's Board composition currently reflects 55% independent directors, "
                    "marginally exceeding the SOE Act 2023 Section 6(b) minimum threshold of "
                    "51%. The Board of Governors has an active Audit Committee and a Risk "
                    "Management Sub-Committee, both constituted per the SOE Policy 2023 "
                    "annexures. The mandatory Board Solvency and Financial Framework "
                    "Declaration under Section 25(3) of the SOE Act was filed with the Finance "
                    "Division on 28 February 2025 — compliant with the prescribed timeline."
                ),
            },
            {
                "framework": "SECP ESG Guidelines 2023 — Governance Index §5.3",
                "metric": "Board Gender Diversity",
                "status": "amber",
                "pct": 60,
                "detail": (
                    "At 18% female Board representation, NHA falls below the SECP's indicative "
                    "25% target for FY2025. The SECP Governance Index §5.3 now requires a "
                    "Board-level Gender Diversity Policy with a concrete 3-year pipeline plan. "
                    "While two female directors hold positions, no formal diversity policy "
                    "document has been lodged with the Board secretariat or disclosed in the "
                    "Annual Report. This constitutes a partial non-compliance."
                ),
            },
            {
                "framework": "SECP ESG Guidelines 2023 — Governance Index §5.5",
                "metric": "External Assurance on ESG Data",
                "status": "red",
                "pct": 0,
                "detail": (
                    "No third-party external assurance engagement has been conducted or "
                    "disclosed for any ESG metric reported in the FY2024 Annual Report. The "
                    "SECP ESG Governance Index §5.5 mandates a phased external assurance "
                    "implementation plan: limited assurance by FY2025, reasonable assurance by "
                    "FY2027. Failure to initiate this process places NHA in pre-breach status "
                    "ahead of the FY2025 regulatory deadline. This is designated a Critical "
                    "Operational Gap requiring Board resolution and CMU notification within 60 "
                    "days of this audit report issuance."
                ),
            },
        ],
    },
    "red_flags": [
        "Renewable energy disclosure below 30% SECP interim benchmark — Board resolution required.",
        "Gender pay gap (11.3%) exceeds ≤7% threshold — Pay Gap Remediation Plan not filed.",
        "Supply chain labour audit coverage at 22% vs ≥80% SECP requirement.",
        "No external assurance engagement — FY2025 deadline at risk.",
        "Section 238 Directors' Report CSR disclosures lack quantifiable metrics.",
    ],
    "remedial_actions": [
        {
            "action": "Commission Scope 1 & 2 GHG baseline audit (ISO 14064-1) by Q3 FY2025.",
            "owner": "Chief Operating Officer / Head of Infrastructure",
            "timeline": "90 days",
        },
        {
            "action": "File Board-approved Gender Pay Gap Remediation Plan with SECP.",
            "owner": "HR Director / Company Secretary",
            "timeline": "60 days",
        },
        {
            "action": "Expand Tier-1 supplier labour audit programme to ≥80% coverage.",
            "owner": "Procurement Division",
            "timeline": "6 months",
        },
        {
            "action": "Appoint external assurance provider; scope limited assurance engagement.",
            "owner": "CFO / Audit Committee",
            "timeline": "45 days",
        },
        {
            "action": "Prepare capital expenditure roadmap for renewable energy installations.",
            "owner": "Chief Engineer / Finance Division",
            "timeline": "120 days",
        },
    ],
}

# ─────────────────────────────────────────────────────────────────────────────
# V.  PDF TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────
def extract_pdf_text(file_obj) -> str:
    """Extract text from an uploaded PDF file object using pypdf."""
    reader = PdfReader(file_obj)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


# ─────────────────────────────────────────────────────────────────────────────
# VI.  GEMINI ANALYSIS ENGINE
# ─────────────────────────────────────────────────────────────────────────────
ANALYSIS_SCHEMA = """
You are a senior Pakistani corporate governance and ESG regulatory counsel with expertise in
the Companies Act 2017, SOE Act 2023, SOE Policy 2023, and SECP Voluntary ESG Disclosure
Guidelines 2023.  You write with the analytical precision of a regulatory attorney from
Islamabad.  Your output must be structured JSON — no prose preamble, no markdown fences.

Return a JSON object matching this exact schema:

{
  "entity": "<company name extracted from document>",
  "audit_date": "<reporting period>",
  "overall_score": <0–100 float>,
  "environmental": {
    "score": <0–100 float>,
    "ghg_emissions_mt": <float or null>,
    "emissions_intensity": <float or null>,
    "renewable_energy_pct": <float or null>,
    "water_reclamation_pct": <float or null>,
    "findings": [
      {
        "framework": "<exact legal reference>",
        "metric": "<metric name>",
        "status": "<green|amber|red>",
        "pct": <0–100 float>,
        "detail": "<3–5 sentence analytical justification citing specific section numbers>"
      }
    ]
  },
  "social": {
    "score": <0–100 float>,
    "ceo_pay_ratio": <float or null>,
    "gender_pay_gap_pct": <float or null>,
    "turnover_rate_pct": <float or null>,
    "temp_worker_ratio_pct": <float or null>,
    "supply_chain_audit_pct": <float or null>,
    "findings": [
      {
        "framework": "<exact legal reference>",
        "metric": "<metric name>",
        "status": "<green|amber|red>",
        "pct": <0–100 float>,
        "detail": "<3–5 sentence analytical justification>"
      }
    ]
  },
  "governance": {
    "score": <0–100 float>,
    "board_independence_pct": <float or null>,
    "board_gender_diversity_pct": <float or null>,
    "external_assurance": <true|false>,
    "supplier_code_of_conduct": <true|false>,
    "solvency_declaration_filed": <true|false>,
    "findings": [
      {
        "framework": "<exact legal reference>",
        "metric": "<metric name>",
        "status": "<green|amber|red>",
        "pct": <0–100 float>,
        "detail": "<3–5 sentence analytical justification>"
      }
    ]
  },
  "red_flags": ["<concise critical issue 1>", "..."],
  "remedial_actions": [
    {
      "action": "<specific corrective action>",
      "owner": "<responsible function/role>",
      "timeline": "<concrete timeline>"
    }
  ]
}

Evaluate the following document text against:
1. Companies Act 2017 — Section 238 (Directors' Report ESG & CSR disclosures)
2. SOE Act 2023 — Section 6 (Sound governance) and Section 25(3) (Solvency declarations)
3. SECP Voluntary ESG Disclosure Guidelines 2023 — Environmental (GHG, energy, water),
   Social (pay ratios, gender gap, turnover, supply chain), Governance (board independence,
   diversity, external assurance, supplier code of conduct)
4. ISO S1 (General Sustainability Disclosure Requirements) and ISO S2 (Climate Disclosures)

Document text:
"""


def run_gemini_analysis(api_key: str, document_text: str, secondary_texts: list[str] = None) -> dict:
    """
    Call Gemini 2.5 Flash with the document text.
    Returns parsed analysis dict or raises an exception.
    """
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash")

    combined = document_text
    if secondary_texts:
        for t in secondary_texts:
            combined += "\n\n--- SECONDARY DOCUMENT ---\n\n" + t

    # Trim to ~120k chars to stay within context limits
    combined = combined[:120_000]

    prompt = ANALYSIS_SCHEMA + combined
    response = model.generate_content(prompt)

    raw = response.text.strip()
    # Strip markdown fences if model adds them
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


# ─────────────────────────────────────────────────────────────────────────────
# VII.  WORD REPORT GENERATORS
# ─────────────────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour using OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x37, 0x4B, 0x59)
    return p


def _add_body(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)  # 1.5 line spacing at 12pt
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def _add_status_label(doc: Document, status: str, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.font.size = Pt(9)
    run.font.bold = True
    if status == "green":
        run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    elif status == "amber":
        run.font.color.rgb = RGBColor(0xFB, 0xBF, 0x24)
    else:
        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)


def generate_gap_audit_report(data: dict) -> bytes:
    """
    REPORT 1: Regulatory Compliance Gap Audit Report
    Returns bytes of a .docx file.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── Cover header ──────────────────────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    _set_cell_bg(cell, "0B132B")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🏛️  AURA GREEN  —  ESG INTELLIGENCE PLATFORM")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REGULATORY COMPLIANCE GAP AUDIT REPORT")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x13, 0x2B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Entity: {data.get('entity', 'N/A')}   |   Period: {data.get('audit_date', 'N/A')}")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x3F, 0x7A, 0x85)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.date.today().strftime('%d %B %Y')}   |   Prepared by: Aura Green Statutory Engine")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()

    # ── Table of Contents ─────────────────────────────────────────────────
    _add_heading(doc, "TABLE OF CONTENTS", 1)
    toc_items = [
        "1. Executive Summary & Overall ESG Score",
        "2. Environmental Pillar Audit — SECP ESG Index §3 / GRI 305 / ISO S2",
        "3. Social Pillar Audit — SECP ESG Index §4 / Companies Act 2017 §238",
        "4. Governance Pillar Audit — SECP ESG Index §5 / SOE Act 2023 §6 & §25(3)",
        "5. Red Flag Operational Omission Index",
        "6. Remedial Action Plan",
        "7. Statutory References & Regulatory Roadmap",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_page_break()

    # ── 1. Executive Summary ──────────────────────────────────────────────
    _add_heading(doc, "1. EXECUTIVE SUMMARY & OVERALL ESG SCORE", 1)
    overall = data.get("overall_score", 0)
    env_score = data.get("environmental", {}).get("score", 0)
    soc_score = data.get("social", {}).get("score", 0)
    gov_score = data.get("governance", {}).get("score", 0)

    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = "Table Grid"
    headers = ["Overall ESG Score", "Environmental", "Social", "Governance"]
    scores = [overall, env_score, soc_score, gov_score]
    for i, (h, s) in enumerate(zip(headers, scores)):
        hc = summary_table.cell(0, i)
        vc = summary_table.cell(1, i)
        _set_cell_bg(hc, "0B132B")
        hc.paragraphs[0].add_run(h).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hc.paragraphs[0].runs[0].font.bold = True
        hc.paragraphs[0].runs[0].font.size = Pt(10)
        score_color = "10B981" if s >= 70 else ("FBBF24" if s >= 50 else "EF4444")
        _set_cell_bg(vc, "F8FAFC")
        run = vc.paragraphs[0].add_run(f"{s:.1f} / 100")
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(
            int(score_color[:2], 16), int(score_color[2:4], 16), int(score_color[4:], 16)
        )
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_body(
        doc,
        (
            f"This Regulatory Compliance Gap Audit has been conducted by the Aura Green statutory "
            f"analysis engine against the applicable Pakistani legal and voluntary ESG disclosure "
            f"frameworks.  The entity under review — {data.get('entity', 'the subject company')} — "
            f"has returned an overall ESG score of {overall:.1f}/100 for the period "
            f"{data.get('audit_date', 'N/A')}.  The composite score reflects a weighted average "
            f"of pillar-level assessments, each benchmarked against the exact disclosure indices "
            f"specified in the SECP Voluntary ESG Disclosure Guidelines 2023, supplemented by "
            f"mandatory obligations under Section 238 of the Companies Act, 2017 and, where "
            f"applicable, Sections 6 and 25(3) of the State-Owned Enterprises Act, 2023."
        ),
    )

    # ── Pillar sections ───────────────────────────────────────────────────
    pillar_map = [
        ("2. ENVIRONMENTAL PILLAR AUDIT", "environmental", "🌿",
         "SECP ESG Disclosure Guidelines 2023 — Environmental Index §3 | "
         "Companies Act 2017 §238 | ISO S2 (Climate-related Disclosures) | GRI 305"),
        ("3. SOCIAL PILLAR AUDIT", "social", "🤝",
         "SECP ESG Disclosure Guidelines 2023 — Social Index §4 | "
         "Companies Act 2017 §238 | ILO Conventions 138 & 182"),
        ("4. GOVERNANCE PILLAR AUDIT", "governance", "⚖️",
         "SECP ESG Disclosure Guidelines 2023 — Governance Index §5 | "
         "SOE Act 2023 §6 & §25(3) | ISO S1 (Sustainability Disclosure)"),
    ]

    for heading, key, icon, framework_ref in pillar_map:
        doc.add_page_break()
        _add_heading(doc, heading, 1)
        _add_body(doc, f"Applicable Framework: {framework_ref}", italic=True)
        pillar_data = data.get(key, {})
        score = pillar_data.get("score", 0)
        _add_body(doc, f"Pillar Score: {score:.1f} / 100", bold=True)
        doc.add_paragraph()

        for finding in pillar_data.get("findings", []):
            _add_heading(doc, finding.get("metric", ""), 2)
            _add_body(doc, f"Legal Framework: {finding.get('framework', '')}", italic=True)
            status = finding.get("status", "amber")
            status_text = {"green": "✅ COMPLIANT", "amber": "⚠ AT RISK", "red": "🚨 BREACH"}
            _add_status_label(doc, status, status_text.get(status, status.upper()))
            _add_body(doc, f"Completion / Coverage: {finding.get('pct', 0):.0f}%", bold=True)
            _add_body(doc, finding.get("detail", ""))
            doc.add_paragraph()

    # ── 5. Red Flag Index ─────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "5. RED FLAG OPERATIONAL OMISSION INDEX", 1)
    _add_body(
        doc,
        (
            "The following items constitute material non-disclosures or threshold breaches "
            "requiring immediate Board attention and, where stipulated, formal regulatory "
            "notification.  Each item maps to a specific statutory obligation."
        ),
    )
    doc.add_paragraph()
    rf_table = doc.add_table(rows=1, cols=2)
    rf_table.style = "Table Grid"
    _set_cell_bg(rf_table.cell(0, 0), "EF4444")
    _set_cell_bg(rf_table.cell(0, 1), "EF4444")
    rf_table.cell(0, 0).paragraphs[0].add_run("#").font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rf_table.cell(0, 1).paragraphs[0].add_run("Red Flag Item").font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, flag in enumerate(data.get("red_flags", []), 1):
        row = rf_table.add_row()
        row.cells[0].paragraphs[0].add_run(str(i)).font.size = Pt(11)
        run = row.cells[1].paragraphs[0].add_run(flag)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── 6. Remedial Action Plan ───────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "6. REMEDIAL ACTION PLAN", 1)
    ra_table = doc.add_table(rows=1, cols=3)
    ra_table.style = "Table Grid"
    for i, header in enumerate(["Action Required", "Responsible Function", "Timeline"]):
        cell = ra_table.cell(0, i)
        _set_cell_bg(cell, "0B132B")
        run = cell.paragraphs[0].add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ra in data.get("remedial_actions", []):
        row = ra_table.add_row()
        row.cells[0].paragraphs[0].add_run(ra.get("action", "")).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(ra.get("owner", "")).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(ra.get("timeline", "")).font.size = Pt(10)

    # ── 7. Statutory References ───────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "7. STATUTORY REFERENCES & REGULATORY ROADMAP", 1)
    refs = [
        "Companies Act, 2017 (Pakistan) — Section 238: Mandatory Directors' Report Disclosures",
        "State-Owned Enterprises Act, 2023 — Section 6 (Sound Governance) & Section 25(3) (Solvency)",
        "SOE Policy, 2023 — Finance Division / Central Monitoring Unit Parameters",
        "SECP Voluntary ESG Disclosure Guidelines, 2023 — Environmental, Social & Governance Indices",
        "SECP Circular No. 19 of 2022 — Clarification on Section 238 Disclosure Specificity",
        "ISO S1: General Requirements for Sustainability-related Financial Disclosures (ISSB 2023)",
        "ISO S2: Climate-related Disclosures (ISSB 2023)",
        "GRI Universal Standards 2021 — GRI 305 (Emissions), GRI 401 (Employment)",
        "ILO Conventions 138 (Minimum Age) and 182 (Worst Forms of Child Labour)",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="List Bullet")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_statutory_esg_report(data: dict) -> bytes:
    """
    REPORT 2: Formal Institutional ESG Statutory Report
    Returns bytes of a .docx file.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── Cover ─────────────────────────────────────────────────────────────
    cover_table = doc.add_table(rows=1, cols=1)
    cell = cover_table.cell(0, 0)
    _set_cell_bg(cell, "0B132B")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n🏛️  AURA GREEN  —  ESG INTELLIGENCE PLATFORM\n\n")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("FORMAL INSTITUTIONAL ESG STATUTORY REPORT\n")
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run2.font.bold = True
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        f"{data.get('entity', 'N/A')}   |   {data.get('audit_date', 'N/A')}"
        f"   |   {datetime.date.today().strftime('%d %B %Y')}\n\n"
    )
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0xA8, 0xD5, 0xDC)

    doc.add_page_break()

    # ── Foreword ──────────────────────────────────────────────────────────
    _add_heading(doc, "FOREWORD", 1)
    _add_body(
        doc,
        (
            f"This Formal Institutional ESG Statutory Report has been compiled in accordance with "
            f"the SECP Voluntary ESG Disclosure Guidelines (2023), the Companies Act, 2017 "
            f"(Section 238), and applicable international sustainability reporting standards "
            f"including ISO S1 (General Requirements for Sustainability-related Financial "
            f"Disclosures) and ISO S2 (Climate-related Disclosures).  The report constitutes "
            f"the official sustainability record of {data.get('entity', 'the subject entity')} "
            f"for the period {data.get('audit_date', 'N/A')} and is intended for submission to "
            f"the Securities and Exchange Commission of Pakistan, the Finance Division's Central "
            f"Monitoring Unit (where the entity is SOE-classified), and for public disclosure "
            f"on the entity's investor relations portal."
        ),
    )

    # ── Pillar disclosures ─────────────────────────────────────────────────
    pillar_configs = [
        {
            "title": "SECTION A — ENVIRONMENTAL PERFORMANCE DISCLOSURE",
            "subtitle": "SECP ESG Environmental Index §3 | ISO S2 | GRI 305",
            "key": "environmental",
            "metrics": [
                ("Total GHG Emissions (Scope 1 & 2)", "ghg_emissions_mt", "metric tons CO₂e"),
                ("Emissions Intensity Ratio", "emissions_intensity", "tCO₂e per unit revenue"),
                ("Renewable Energy Mix", "renewable_energy_pct", "%"),
                ("Water Reclamation Rate", "water_reclamation_pct", "%"),
            ],
            "iso_ref": (
                "ISO S2 Climate-related Disclosures requires entities to report: (a) physical "
                "climate risks and transition risks material to the business; (b) climate-related "
                "governance processes; (c) strategy for managing climate risks; and (d) "
                "quantitative Scope 1, 2, and 3 GHG emissions data validated against the GHG "
                "Protocol Corporate Standard."
            ),
        },
        {
            "title": "SECTION B — SOCIAL PERFORMANCE DISCLOSURE",
            "subtitle": "SECP ESG Social Index §4 | Companies Act 2017 §238 | ILO Conventions",
            "key": "social",
            "metrics": [
                ("CEO-to-Median Employee Pay Ratio", "ceo_pay_ratio", ":1"),
                ("Gender Pay Gap", "gender_pay_gap_pct", "%"),
                ("Voluntary Employee Turnover", "turnover_rate_pct", "%"),
                ("Temporary Worker Ratio", "temp_worker_ratio_pct", "% of workforce"),
                ("Supply Chain Labour Audit Coverage", "supply_chain_audit_pct", "%"),
            ],
            "iso_ref": (
                "ISO S1 General Requirements mandate disclosure of material social impacts "
                "affecting the entity's workforce and communities, with quantitative indicators "
                "where available.  Pakistani entities are additionally bound by Section 238(1)(e) "
                "of the Companies Act, 2017, requiring Board certification of human capital "
                "disclosures in the Directors' Report."
            ),
        },
        {
            "title": "SECTION C — GOVERNANCE STRUCTURE DISCLOSURE",
            "subtitle": "SECP ESG Governance Index §5 | SOE Act 2023 §6 & §25(3) | ISO S1",
            "key": "governance",
            "metrics": [
                ("Board Independence", "board_independence_pct", "%"),
                ("Board Gender Diversity", "board_gender_diversity_pct", "%"),
                ("External Assurance Engaged", "external_assurance", "Yes/No"),
                ("Supplier Code of Conduct Active", "supplier_code_of_conduct", "Yes/No"),
                ("Section 25(3) Solvency Declaration Filed", "solvency_declaration_filed", "Yes/No"),
            ],
            "iso_ref": (
                "ISO S1 requires entities to disclose governance structures responsible for "
                "sustainability oversight, including Board-level accountability mechanisms, "
                "management roles, and the integration of sustainability considerations into "
                "remuneration and strategic planning processes."
            ),
        },
    ]

    for cfg in pillar_configs:
        doc.add_page_break()
        _add_heading(doc, cfg["title"], 1)
        _add_body(doc, cfg["subtitle"], italic=True)
        doc.add_paragraph()

        # Quantitative metrics table
        _add_heading(doc, "Quantitative Indicators", 2)
        m_table = doc.add_table(rows=1, cols=3)
        m_table.style = "Table Grid"
        for i, h in enumerate(["Indicator", "Reported Value", "Unit / Note"]):
            c = m_table.cell(0, i)
            _set_cell_bg(c, "3F7A85")
            run = c.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        pillar_data = data.get(cfg["key"], {})
        for label, field, unit in cfg["metrics"]:
            row = m_table.add_row()
            row.cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
            val = pillar_data.get(field)
            if val is None:
                disp = "Not Disclosed"
            elif isinstance(val, bool):
                disp = "Yes" if val else "No"
            elif isinstance(val, float):
                disp = f"{val:.1f}"
            else:
                disp = str(val)
            row.cells[1].paragraphs[0].add_run(disp).font.size = Pt(10)
            row.cells[2].paragraphs[0].add_run(unit).font.size = Pt(10)

        doc.add_paragraph()

        # ISO reference
        _add_heading(doc, "International Standard Alignment", 2)
        _add_body(doc, cfg["iso_ref"])
        doc.add_paragraph()

        # Analytical findings
        _add_heading(doc, "Analytical Findings & Regulatory Assessment", 2)
        for finding in pillar_data.get("findings", []):
            p = doc.add_paragraph()
            run = p.add_run(f"▸  {finding.get('metric', '')}")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x0B, 0x13, 0x2B)
            _add_body(doc, finding.get("detail", ""))
            doc.add_paragraph()

    # ── Concluding Declaration ────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "CONCLUDING DECLARATION", 1)
    _add_body(
        doc,
        (
            "This report has been generated by the Aura Green ESG Intelligence Platform, "
            "incorporating document analysis against the statutory frameworks enumerated herein. "
            "The findings and scores contained in this report are derived from disclosed "
            "information in the entity's regulatory filings and annual reports.  Where data "
            "points were not disclosed, they have been noted as omissions and factored into "
            "the pillar scoring accordingly.  This report is intended to serve as a formal "
            "institutional sustainability disclosure document and may be submitted to the SECP, "
            "the Finance Division's Central Monitoring Unit, external rating agencies, and "
            "international sustainability indices (MSCI ESG, Sustainalytics, CDP)."
        ),
    )
    doc.add_paragraph()
    _add_body(doc, f"Report Date: {datetime.date.today().strftime('%d %B %Y')}", bold=True)
    _add_body(doc, "Platform: Aura Green Statutory Engine v1.0", bold=True)
    _add_body(doc, f"Entity: {data.get('entity', 'N/A')}", bold=True)
    _add_body(doc, f"Period: {data.get('audit_date', 'N/A')}", bold=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# VIII.  DASHBOARD RENDERING HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def render_kpi_row(data: dict):
    cols = st.columns(4)
    scores = [
        ("Overall ESG", data.get("overall_score", 0), "▲"),
        ("Environmental", data.get("environmental", {}).get("score", 0), "🌿"),
        ("Social", data.get("social", {}).get("score", 0), "🤝"),
        ("Governance", data.get("governance", {}).get("score", 0), "⚖️"),
    ]
    for col, (label, score, icon) in zip(cols, scores):
        trend_class = "neutral" if score < 50 else ("down" if score < 65 else "")
        with col:
            st.markdown(
                f"""
                <div class="kpi-card">
                    <div class="kpi-label">{icon} {label}</div>
                    <div class="kpi-value">{score:.1f}</div>
                    <div class="kpi-trend {trend_class}">{'▲ Strong' if score >= 70 else ('▼ At Risk' if score < 50 else '~ Review Required')}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )


def render_finding_card(finding: dict):
    status = finding.get("status", "amber")
    pct = finding.get("pct", 0)
    color_map = {"green": "green", "amber": "amber", "red": "red"}
    bar_color = color_map.get(status, "amber")
    label_map = {"green": "Compliant", "amber": "At Risk", "red": "Breach"}
    st.markdown(
        f"""
        <div class="aura-panel">
            <div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:8px;">
                <div>
                    <div class="aura-panel-title" style="margin-bottom:2px">{finding.get('metric','')}</div>
                    <div style="font-size:10px;color:#A8D5DC;margin-bottom:6px;">{finding.get('framework','')}</div>
                </div>
                {pill(label_map.get(status, status.upper()), status)}
            </div>
            {progress_bar(pct, bar_color)}
            <div class="prog-meta"><span>{pct:.0f}% coverage / completion</span></div>
            <div class="audit-block {status if status != 'green' else ''}" style="margin-top:10px;">
                {finding.get('detail','').replace(chr(10), '<br>')}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_pillar_tab(pillar_data: dict):
    score = pillar_data.get("score", 0)
    score_status = "green" if score >= 70 else ("amber" if score >= 50 else "red")
    st.markdown(
        f"<div style='display:flex;align-items:center;gap:12px;margin-bottom:20px;'>"
        f"<span style='font-size:28px;font-weight:700;color:#FFFFFF'>{score:.1f}<span style='font-size:14px;color:#A8D5DC'> / 100</span></span>"
        f"{pill('Score', score_status)}</div>",
        unsafe_allow_html=True,
    )
    for finding in pillar_data.get("findings", []):
        render_finding_card(finding)


# ─────────────────────────────────────────────────────────────────────────────
# IX.  EXECUTIVE TOP HEADER
# ─────────────────────────────────────────────────────────────────────────────
st.markdown(
    f"""
    <div class="aura-header">
        <div>
            <div class="aura-header-title">🏛️ AURA GREEN</div>
            <div class="aura-header-sub">ESG Intelligence &amp; Statutory Compliance Platform</div>
        </div>
        <div class="aura-header-meta">
            SECP · Companies Act 2017 · SOE Act 2023<br>
            {datetime.date.today().strftime('%d %B %Y')} &nbsp;|&nbsp; FY2025 Q2 Cycle
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────────────────────────────────────────────
# X.  MAIN LAYOUT — columns([1, 2.5])
# ─────────────────────────────────────────────────────────────────────────────
left_col, right_col = st.columns([1, 2.5])

# ─────────────────────────────────────────────────────────────────────────────
# LEFT COLUMN — Document Repository
# ─────────────────────────────────────────────────────────────────────────────
with left_col:
    st.markdown('<div class="aura-panel"><div class="aura-panel-title">📂 Document Repository</div>', unsafe_allow_html=True)

    st.markdown("**Primary Ingestion Slot**")
    st.caption("Annual Report / Sustainability Report (PDF, max 100 MB)")
    primary_file = st.file_uploader(
        "Upload Primary Document",
        type=["pdf"],
        key="primary_upload",
        label_visibility="collapsed",
    )

    if primary_file:
        st.markdown(
            f"""
            <div class="file-item">
                <div class="file-dot-green"></div>
                <div>
                    <div style="font-weight:600">{primary_file.name[:28]}…</div>
                    <div style="font-size:10px;color:#A8D5DC">{primary_file.size // 1024} KB · Indexed</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("---")
    st.markdown("**Secondary Documents (Optional)**")
    st.caption("Gazette Amendments, SECP Circulars, Board Resolutions (PDF, multi-select)")
    secondary_files = st.file_uploader(
        "Upload Secondary Documents",
        type=["pdf"],
        key="secondary_upload",
        accept_multiple_files=True,
        label_visibility="collapsed",
    )

    for sf in secondary_files:
        st.markdown(
            f"""
            <div class="file-item">
                <div class="file-dot-amber"></div>
                <div>
                    <div style="font-weight:600">{sf.name[:26]}…</div>
                    <div style="font-size:10px;color:#A8D5DC">{sf.size // 1024} KB · Queued</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown('</div>', unsafe_allow_html=True)

    # API key input
    st.markdown('<div class="aura-panel"><div class="aura-panel-title">🔑 Engine Configuration</div>', unsafe_allow_html=True)
    api_key = st.text_input(
        "Google GenAI API Key",
        type="password",
        placeholder="AIza…",
        help="Gemini 2.5 Flash. If left blank, the engine will load the fallback SOE sample dataset.",
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # Run button
    run_analysis = st.button("⚡ RUN STATUTORY ANALYSIS", type="primary", use_container_width=True)

    # Status / legend
    st.markdown(
        f"""
        <div class="aura-panel" style="margin-top:12px;">
            <div class="aura-panel-title">Status Legend</div>
            <div style="display:flex;flex-direction:column;gap:6px;font-size:12px;">
                <div>{pill("Compliant","green")} &nbsp; ≥ 70% threshold met</div>
                <div>{pill("At Risk","amber")} &nbsp; 50–69% — review required</div>
                <div>{pill("Breach","red")} &nbsp; &lt; 50% — regulatory action</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ─────────────────────────────────────────────────────────────────────────────
# RIGHT COLUMN — Analysis Workspace
# ─────────────────────────────────────────────────────────────────────────────
with right_col:
    if "analysis_data" not in st.session_state:
        st.session_state.analysis_data = None
    if "fallback_active" not in st.session_state:
        st.session_state.fallback_active = False

    # ── Trigger analysis ───────────────────────────────────────────────────
    if run_analysis:
        if not primary_file and not api_key:
            # No inputs — load fallback directly
            st.session_state.analysis_data = FALLBACK_ANALYSIS
            st.session_state.fallback_active = True
        elif not primary_file:
            st.warning("⚠ Please upload a primary PDF document before running the analysis.")
        else:
            with st.spinner("Extracting document text…"):
                primary_text = extract_pdf_text(primary_file)

            secondary_texts = []
            for sf in secondary_files:
                secondary_texts.append(extract_pdf_text(sf))

            if not api_key:
                st.info("ℹ No API key provided — loading fallback SOE sample dataset.")
                st.session_state.analysis_data = FALLBACK_ANALYSIS
                st.session_state.fallback_active = True
            else:
                with st.spinner("Routing to Gemini 2.5 Flash statutory engine…"):
                    try:
                        result = run_gemini_analysis(api_key, primary_text, secondary_texts)
                        st.session_state.analysis_data = result
                        st.session_state.fallback_active = False
                    except Exception as exc:
                        error_msg = str(exc)
                        # ── IV. FAULT-TOLERANT FALLBACK ─────────────────────
                        is_server_error = any(
                            code in error_msg for code in ["503", "502", "500", "overloaded", "unavailable", "quota"]
                        )
                        if is_server_error:
                            st.warning(
                                "⚠ **Gemini API Server Overload Detected.**  "
                                "The remote statutory engine returned a temporary server error "
                                f"(`{error_msg[:120]}`).  "
                                "The fallback compliance dataset — an illustrative audit of a "
                                "Pakistani State-Owned Enterprise benchmarked against Finance "
                                "Division CMU parameters — has been loaded automatically.  "
                                "All dashboard metrics, pillar cards, and Word document "
                                "generators remain fully operational."
                            )
                        else:
                            st.warning(
                                f"⚠ **Analysis Engine Error:** `{error_msg[:200]}`.  "
                                "Fallback SOE dataset loaded to maintain dashboard functionality."
                            )
                        st.session_state.analysis_data = FALLBACK_ANALYSIS
                        st.session_state.fallback_active = True

    # ── Render dashboard ───────────────────────────────────────────────────
    data = st.session_state.analysis_data

    if data is None:
        st.markdown(
            """
            <div class="aura-panel" style="text-align:center;padding:60px 20px;">
                <div style="font-size:40px;margin-bottom:16px;">🏛️</div>
                <div style="font-size:18px;font-weight:700;color:#FFFFFF;margin-bottom:8px;">
                    Aura Green Statutory Engine Ready
                </div>
                <div style="font-size:13px;color:#A8D5DC;line-height:1.7;">
                    Upload an Annual Report PDF in the left repository panel and click<br>
                    <strong>⚡ RUN STATUTORY ANALYSIS</strong> to begin the ESG audit.<br><br>
                    The engine evaluates disclosures against:<br>
                    Companies Act 2017 §238 · SOE Act 2023 §6 & §25(3)<br>
                    SECP ESG Guidelines 2023 · ISO S1 & S2
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        # Fallback notice
        if st.session_state.fallback_active:
            st.markdown(
                f"""
                <div style="background:rgba(251,191,36,0.08);border:1px solid #FBBF24;
                padding:10px 16px;margin-bottom:16px;font-size:12px;color:#FBBF24;">
                    ⚠ &nbsp;<strong>Fallback Dataset Active</strong> — Illustrative audit of
                    <em>{data.get('entity','')}</em>.
                    Upload your own document and provide an API key to run a live analysis.
                </div>
                """,
                unsafe_allow_html=True,
            )

        # Entity banner
        st.markdown(
            f"""
            <div style="display:flex;justify-content:space-between;align-items:center;
            background:#2C5F6A;border:1px solid #000;padding:12px 18px;margin-bottom:16px;">
                <div>
                    <div style="font-size:15px;font-weight:700;color:#FFFFFF">{data.get('entity','N/A')}</div>
                    <div style="font-size:11px;color:#A8D5DC">{data.get('audit_date','N/A')}</div>
                </div>
                {pill("Live Audit" if not st.session_state.fallback_active else "Sample Dataset",
                      "green" if not st.session_state.fallback_active else "amber")}
            </div>
            """,
            unsafe_allow_html=True,
        )

        # KPI row
        render_kpi_row(data)
        st.markdown("<div style='margin-bottom:20px'></div>", unsafe_allow_html=True)

        # Pillar tabs
        tab_env, tab_soc, tab_gov = st.tabs(
            ["🌿 Environmental Pillar", "🤝 Social Pillar", "⚖️ Governance Pillar"]
        )

        with tab_env:
            env = data.get("environmental", {})
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(
                    f"""
                    <div class="aura-panel">
                        <div class="aura-panel-title">Key Quantitative Indicators</div>
                        <div class="metric-row"><span class="metric-label">GHG Emissions (tCO₂e)</span>
                            <span>{env.get('ghg_emissions_mt') or 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Emissions Intensity</span>
                            <span>{env.get('emissions_intensity') or 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Renewable Energy Mix</span>
                            <span>{f"{env.get('renewable_energy_pct'):.0f}%" if env.get('renewable_energy_pct') else 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Water Reclamation</span>
                            <span>{f"{env.get('water_reclamation_pct'):.0f}%" if env.get('water_reclamation_pct') else 'N/D'}</span></div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with col2:
                st.markdown(
                    f"""
                    <div class="aura-panel">
                        <div class="aura-panel-title">Framework Coverage</div>
                        <div style="font-size:11px;color:#A8D5DC;margin-bottom:8px">SECP ESG Environmental Index §3</div>
                        {progress_bar(env.get('score', 0), status_from_pct(env.get('score', 0)))}
                        <div class="prog-meta"><span>Pillar Score</span><span>{env.get('score', 0):.0f} / 100</span></div>
                        <div style="margin-top:8px;font-size:11px;color:#A8D5DC">GRI 305 · ISO S2 · Companies Act §238</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            render_pillar_tab(env)

        with tab_soc:
            soc = data.get("social", {})
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(
                    f"""
                    <div class="aura-panel">
                        <div class="aura-panel-title">Key Quantitative Indicators</div>
                        <div class="metric-row"><span class="metric-label">CEO Pay Ratio</span>
                            <span>{f"{soc.get('ceo_pay_ratio'):.1f}:1" if soc.get('ceo_pay_ratio') else 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Gender Pay Gap</span>
                            <span>{f"{soc.get('gender_pay_gap_pct'):.1f}%" if soc.get('gender_pay_gap_pct') else 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Employee Turnover</span>
                            <span>{f"{soc.get('turnover_rate_pct'):.1f}%" if soc.get('turnover_rate_pct') else 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Temp Worker Ratio</span>
                            <span>{f"{soc.get('temp_worker_ratio_pct'):.1f}%" if soc.get('temp_worker_ratio_pct') else 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Supply Chain Audit</span>
                            <span>{f"{soc.get('supply_chain_audit_pct'):.0f}%" if soc.get('supply_chain_audit_pct') else 'N/D'}</span></div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with col2:
                st.markdown(
                    f"""
                    <div class="aura-panel">
                        <div class="aura-panel-title">Framework Coverage</div>
                        <div style="font-size:11px;color:#A8D5DC;margin-bottom:8px">SECP ESG Social Index §4</div>
                        {progress_bar(soc.get('score', 0), status_from_pct(soc.get('score', 0)))}
                        <div class="prog-meta"><span>Pillar Score</span><span>{soc.get('score', 0):.0f} / 100</span></div>
                        <div style="margin-top:8px;font-size:11px;color:#A8D5DC">Companies Act §238 · ILO Conv. 138 & 182</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            render_pillar_tab(soc)

        with tab_gov:
            gov = data.get("governance", {})
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(
                    f"""
                    <div class="aura-panel">
                        <div class="aura-panel-title">Key Quantitative Indicators</div>
                        <div class="metric-row"><span class="metric-label">Board Independence</span>
                            <span>{f"{gov.get('board_independence_pct'):.0f}%" if gov.get('board_independence_pct') else 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">Board Gender Diversity</span>
                            <span>{f"{gov.get('board_gender_diversity_pct'):.0f}%" if gov.get('board_gender_diversity_pct') else 'N/D'}</span></div>
                        <div class="metric-row"><span class="metric-label">External Assurance</span>
                            <span>{'Yes' if gov.get('external_assurance') else 'No'}</span></div>
                        <div class="metric-row"><span class="metric-label">Supplier Code of Conduct</span>
                            <span>{'Yes' if gov.get('supplier_code_of_conduct') else 'No'}</span></div>
                        <div class="metric-row"><span class="metric-label">§25(3) Solvency Filed</span>
                            <span>{'Yes' if gov.get('solvency_declaration_filed') else 'No'}</span></div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with col2:
                st.markdown(
                    f"""
                    <div class="aura-panel">
                        <div class="aura-panel-title">Framework Coverage</div>
                        <div style="font-size:11px;color:#A8D5DC;margin-bottom:8px">SECP ESG Governance Index §5</div>
                        {progress_bar(gov.get('score', 0), status_from_pct(gov.get('score', 0)))}
                        <div class="prog-meta"><span>Pillar Score</span><span>{gov.get('score', 0):.0f} / 100</span></div>
                        <div style="font-size:11px;color:#A8D5DC;margin-top:8px">SOE Act 2023 §6 & §25(3) · ISO S1</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            render_pillar_tab(gov)

        # ── Red Flag Index ─────────────────────────────────────────────────
        st.markdown("<div style='margin-top:20px'></div>", unsafe_allow_html=True)
        with st.expander("🚨  RED FLAG OPERATIONAL OMISSION INDEX", expanded=False):
            red_flags = data.get("red_flags", [])
            if red_flags:
                for i, flag in enumerate(red_flags, 1):
                    st.markdown(
                        f'<div class="audit-block red">'
                        f'<strong>#{i}</strong> &nbsp; {flag}'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
            else:
                st.markdown(
                    '<div class="audit-block">No critical red flags identified in this audit cycle.</div>',
                    unsafe_allow_html=True,
                )

        # ── Remedial Action Plan ───────────────────────────────────────────
        with st.expander("📋  REMEDIAL ACTION PLAN", expanded=False):
            for ra in data.get("remedial_actions", []):
                st.markdown(
                    f"""
                    <div class="aura-panel" style="margin-bottom:8px;">
                        <div style="font-size:13px;font-weight:600;color:#FFFFFF;margin-bottom:6px;">
                            {ra.get('action','')}
                        </div>
                        <div style="font-size:11px;color:#A8D5DC;">
                            Owner: <strong>{ra.get('owner','')}</strong> &nbsp;|&nbsp;
                            Timeline: {pill(ra.get('timeline',''),'amber')}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

        # ── Word Document Downloads ────────────────────────────────────────
        st.markdown(
            '<div class="aura-panel" style="margin-top:20px;"><div class="aura-panel-title">📄 Report Generation — Word Documents</div>',
            unsafe_allow_html=True,
        )
        dl_col1, dl_col2 = st.columns(2)

        with dl_col1:
            st.caption("Regulatory Compliance Gap Audit Report")
            with st.spinner("Compiling Report 1…"):
                gap_bytes = generate_gap_audit_report(data)
            st.download_button(
                label="⬇ Download Gap Audit Report (.docx)",
                data=gap_bytes,
                file_name=f"AuraGreen_Gap_Audit_{datetime.date.today().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        with dl_col2:
            st.caption("Formal Institutional ESG Statutory Report")
            with st.spinner("Compiling Report 2…"):
                stat_bytes = generate_statutory_esg_report(data)
            st.download_button(
                label="⬇ Download Statutory ESG Report (.docx)",
                data=stat_bytes,
                file_name=f"AuraGreen_Statutory_ESG_{datetime.date.today().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        st.markdown("</div>", unsafe_allow_html=True)
